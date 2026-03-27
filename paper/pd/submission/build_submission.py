"""Build Brain Stimulation submission assets for the PD manuscript."""

from __future__ import annotations

import argparse
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl.comments import Comment
from openpyxl.utils import get_column_letter

from paper.pd.paths import resolve_project_root, summary_root, summary_table_root

DEFAULT_REFERENCE_METHODS_DOC = Path(
    "/Users/mojackhu/Research/STNSNr/manuscript/presubmit_v0.1/Methods.docx"
)
DEFAULT_SUBMISSION_ROOTNAME = "submission"
COMMENT_PATTERN = re.compile(r"\[\[COMMENT:\s*(.*?)\s*\]\]")
COMMENT_AUTHOR = "Codex"
COMMENT_INITIALS = "CX"


@dataclass(frozen=True)
class PanelSpec:
    """Describe one figure panel asset or placeholder."""

    panel_id: str
    panel_role: str
    status: str
    source_path: str | None
    display_model: str
    band_priority: str
    notes: str
    dual_model_support: str = "n/a"
    dual_model_notes: str = ""
    local_filename: str | None = None
    interval_key: tuple[str, str, str, str, int, int] | None = None


def _submission_root(project_root: Path) -> Path:
    return project_root / DEFAULT_SUBMISSION_ROOTNAME


def _manuscript_root(project_root: Path) -> Path:
    return _submission_root(project_root) / "manuscript"


def _figure_table_root(project_root: Path) -> Path:
    return _submission_root(project_root) / "figure&table"


def _normalize_subject_code(number: Any) -> str:
    try:
        value = int(number)
    except Exception:
        return str(number)
    return f"sub-{value:03d}"


def _normalize_checkmark(value: Any) -> str:
    return "Yes" if str(value).strip() == "✓" else "No"


def normalize_subj_paradigm_frame(frame: pd.DataFrame) -> pd.DataFrame:
    """Normalize the cohort availability workbook into one tidy table."""
    columns = {
        "Subject": "SubjectIndex",
        "ID": "SubjectName",
        "Gait": "Sit",
        "Unnamed: 3": "Stand",
        "Unnamed: 4": "Walk",
        "Unnamed: 5": "Turn",
        "Unnamed: 6": "Cycle",
        "Unnamed: 7": "FoG",
        "Med": "Med",
        "Pain": "Pain",
    }
    data = frame.iloc[1:].copy()
    data = data.rename(columns=columns)
    keep = [v for v in columns.values() if v in data.columns]
    data = data.loc[:, keep].copy()
    data["SubjectIndex"] = data["SubjectIndex"].astype(int)
    data["SubjectCode"] = data["SubjectIndex"].map(_normalize_subject_code)
    for col in ["Sit", "Stand", "Walk", "Turn", "Cycle", "FoG", "Med", "Pain"]:
        if col in data.columns:
            data[col] = data[col].map(_normalize_checkmark)
    ordered = [
        "SubjectIndex",
        "SubjectCode",
        "SubjectName",
        "Sit",
        "Stand",
        "Walk",
        "Turn",
        "Cycle",
        "FoG",
        "Med",
        "Pain",
    ]
    return data.loc[:, [c for c in ordered if c in data.columns]].reset_index(drop=True)


def extract_comment_segments(text: str) -> list[tuple[str, str | None]]:
    """Split text into normal chunks and inline comment payloads."""
    segments: list[tuple[str, str | None]] = []
    cursor = 0
    for match in COMMENT_PATTERN.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor : match.start()], None))
        segments.append(("[?]", match.group(1).strip()))
        cursor = match.end()
    if cursor < len(text):
        segments.append((text[cursor:], None))
    return segments


def scalar_artifact_paths_from_pdf(pdf_path: str | Path) -> tuple[Path, Path] | None:
    """Return `(lmer_csv, rlmer_csv)` for a modeled scalar PDF."""
    pdf = Path(pdf_path)
    if "/lmer_model/" not in pdf.as_posix() or not pdf.name.endswith("_raw.pdf"):
        return None
    lmer_csv = pdf.with_name(pdf.name.replace("_raw.pdf", "_tukey.csv"))
    rlmer_csv = Path(str(lmer_csv).replace("/lmer_model/", "/rlmer_model/"))
    return lmer_csv, rlmer_csv


def _csv_support(csv_path: Path) -> tuple[bool, float | None, int]:
    if not csv_path.exists():
        return False, None, 0
    frame = pd.read_csv(csv_path)
    if "p_tukey" not in frame.columns:
        return False, None, len(frame)
    min_p = float(frame["p_tukey"].min()) if len(frame) else None
    sig_count = int((frame["p_tukey"] < 0.05).sum())
    return sig_count > 0, min_p, sig_count


def evaluate_scalar_dual_support(pdf_path: str | Path) -> tuple[str, str]:
    """Return dual-support summary for one scalar PDF panel."""
    mapped = scalar_artifact_paths_from_pdf(pdf_path)
    if mapped is None:
        return "n/a", "Descriptive panel or non-modeled PDF."
    lmer_csv, rlmer_csv = mapped
    lmer_ok, lmer_min, lmer_n = _csv_support(lmer_csv)
    rlmer_ok, rlmer_min, rlmer_n = _csv_support(rlmer_csv)
    note = (
        f"lmer_min_p={lmer_min}, lmer_sig_contrasts={lmer_n}; "
        f"rlmer_min_p={rlmer_min}, rlmer_sig_contrasts={rlmer_n}"
    )
    return ("yes" if lmer_ok and rlmer_ok else "no"), note


def load_cycle_interval_catalog(project_root: Path) -> pd.DataFrame:
    path = (
        summary_table_root(project_root)
        / "cycle"
        / "interval"
        / "cycle_timepoint_intervals.csv"
    )
    return pd.read_csv(path)


def evaluate_interval_dual_support(
    interval_catalog: pd.DataFrame,
    interval_key: tuple[str, str, str, str, int, int] | None,
) -> tuple[str, str]:
    """Return dual-support summary for one retained cycle interval."""
    if interval_key is None:
        return "n/a", "No interval selector provided."
    metric, band, region, direction, start_pct, end_pct = interval_key
    mask = (
        (interval_catalog["Metric"] == metric)
        & (interval_catalog["Band"] == band)
        & (interval_catalog["Region"] == region)
        & (interval_catalog["direction"] == direction)
        & (interval_catalog["start_pct"] == start_pct)
        & (interval_catalog["end_pct"] == end_pct)
    )
    subset = interval_catalog.loc[mask].copy()
    if subset.empty:
        return "no", "Interval row not found in cycle_timepoint_intervals.csv."
    row = subset.iloc[0]
    supported = bool(row["min_p_lmer"] < 0.05 and row["min_p_rlmer"] < 0.05)
    note = (
        f"span_pct={row['span_pct']}, min_p_lmer={row['min_p_lmer']}, "
        f"min_p_rlmer={row['min_p_rlmer']}, robust_fraction={row['robust_fraction']}"
    )
    return ("yes" if supported else "no"), note


def _existing_figure1_assets(project_root: Path) -> dict[str, str]:
    root = _figure_table_root(project_root)
    return {
        "leaddbs_png": str(root / "leaddbs.png"),
        "workflow_spec_md": str(root / "FigureX_workflow_spec.md"),
    }


def figure_specs(project_root: Path) -> dict[str, list[PanelSpec]]:
    assets = _existing_figure1_assets(project_root)
    return {
        "figure1": [
            PanelSpec("1a", "Study schema across paradigms", "placeholder", None, "descriptive", "secondary-band", "Cohort overview across rest, motor, turn, and cycle."),
            PanelSpec("1b", "Representative localization and pair definition", "copied", assets["leaddbs_png"], "descriptive", "secondary-band", "Reuse existing localization asset for the software-and-cohort overview.", local_filename="1b.png"),
            PanelSpec("1c", "Toolbox master workflow", "placeholder", assets["workflow_spec_md"], "descriptive", "secondary-band", "Summarize import -> localization -> preprocess -> align -> metrics -> export as the primary software flow.", local_filename="1c.md"),
            PanelSpec("1d", "Import detail", "placeholder", None, "descriptive", "secondary-band", "Show supported recording ingestion, annotations, and standardized inputs."),
            PanelSpec("1e", "Localization detail", "placeholder", None, "descriptive", "secondary-band", "Show anatomy-aware channel labels and STN/SNr mapping."),
            PanelSpec("1f", "Preprocess detail", "placeholder", None, "descriptive", "secondary-band", "Show filtering, bad-span handling, and review-ready QC."),
            PanelSpec("1g", "Align detail", "placeholder", None, "descriptive", "secondary-band", "Show condition, state, event, and continuous phase alignment modes."),
            PanelSpec("1h", "Metrics detail", "placeholder", None, "descriptive", "secondary-band", "Show periodic, aperiodic, burst, coherence, wPLI, ciPLV, PSI, and TRGC outputs."),
            PanelSpec("1i", "Export detail", "placeholder", None, "descriptive", "secondary-band", "Show paper-ready CSV/XLSX outputs, figure panel extraction, and stats exports."),
        ],
        "figure2": [
            PanelSpec("2a", "Rest module settings and software entry", "placeholder", None, "descriptive", "secondary-band", "Toolbox panel for rest module configuration and defaults."),
            PanelSpec("2b", "Representative raw trace and QC", "placeholder", None, "descriptive", "secondary-band", "Illustrate raw trace, cleaning, and preprocessing checkpoints for resting data."),
            PanelSpec("2c", "Representative anatomy-guided pair selection", "placeholder", None, "descriptive", "secondary-band", "Show STN-local, SNr-local, and SNr-STN edge definitions used by the rest module."),
            PanelSpec("2d", "Representative OFF vs ON PSD", "placeholder", None, "descriptive", "secondary-band", "Use a layout placeholder for one example OFF/ON PSD comparison."),
            PanelSpec("2e", "Cohort-level resting spectral overview", "copied", "/Users/mojackhu/Research/pd/summary/table/med/periodic/mean-spectral_summary_trans/Phase-Region_raw.pdf", "descriptive", "beta-low", "Rest spectral overview used as cohort-level context.", local_filename="2e.pdf"),
            PanelSpec("2f", "STN beta-low periodic", "copied", "/Users/mojackhu/Research/pd/summary/table/med/periodic/mean-scalar_summary_trans/Beta_low/lmer_model/Phase-Region_raw.pdf", "lmer", "beta-low", "Primary canonical periodic panel for the rest module.", local_filename="2f.pdf"),
            PanelSpec("2g", "STN beta-low burst rate", "copied", "/Users/mojackhu/Research/pd/summary/table/med/burst/rate-scalar_summary_trans/Beta_low/lmer_model/Phase-Region_raw.pdf", "lmer", "beta-low", "Primary burst-rate validation panel.", local_filename="2g.pdf"),
            PanelSpec("2h", "STN beta-low burst occupation", "copied", "/Users/mojackhu/Research/pd/summary/table/med/burst/occupation-scalar_summary_trans/Beta_low/lmer_model/Phase-Region_raw.pdf", "lmer", "beta-low", "Primary burst-occupation validation panel.", local_filename="2h.pdf"),
            PanelSpec("2i", "SNr/STN secondary-band context", "copied", "/Users/mojackhu/Research/pd/summary/table/med/periodic/mean-scalar_summary_trans/Theta/lmer_model/Phase-Region_raw.pdf", "lmer", "secondary-band", "Secondary-band context panel, subordinate to beta-low narrative.", local_filename="2i.pdf"),
            PanelSpec("2j", "SNr-STN coherence Theta", "copied", "/Users/mojackhu/Research/pd/summary/table/med/coherence/mean-scalar_summary_trans/Theta/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Resting synchrony extension panel.", local_filename="2j.pdf"),
            PanelSpec("2k", "SNr->STN TRGC Delta", "copied", "/Users/mojackhu/Research/pd/summary/table/med/trgc/mean-scalar_summary_trans/Delta/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Directional extension panel retained because both models support the result.", local_filename="2k.pdf"),
        ],
        "figure3": [
            PanelSpec("3a", "State module settings", "placeholder", None, "descriptive", "secondary-band", "Toolbox panel for standing-to-walking state analysis."),
            PanelSpec("3b", "State annotation timeline", "placeholder", None, "descriptive", "secondary-band", "Show stand and walk blocks with annotation logic."),
            PanelSpec("3c", "Representative state-aligned raw trace", "placeholder", None, "descriptive", "secondary-band", "Illustrate one raw trace aligned to locomotor states."),
            PanelSpec("3d", "Motor spectral overview", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/periodic/mean-spectral_summary_trans/Phase-Region_raw.pdf", "descriptive", "beta-low", "Cohort-level motor spectral overview.", local_filename="3d.pdf"),
            PanelSpec("3e", "Representative burst detection", "placeholder", None, "descriptive", "secondary-band", "Show how dynamic burst features are extracted in the state module."),
            PanelSpec("3f", "STN beta-low burst rate", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/burst/rate-scalar_summary_trans/Beta_low/lmer_model/Phase-Region_raw.pdf", "lmer", "beta-low", "Primary dynamic beta-low panel.", local_filename="3f.pdf"),
            PanelSpec("3g", "STN beta-low burst occupation", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/burst/occupation-scalar_summary_trans/Beta_low/lmer_model/Phase-Region_raw.pdf", "lmer", "beta-low", "Primary beta-low occupation panel.", local_filename="3g.pdf"),
            PanelSpec("3h", "Beta-high context", "placeholder", None, "descriptive", "beta-high", "Reserve one context panel for beta-high descriptive support if needed in layout."),
            PanelSpec("3i", "Directional edge definition and settings", "placeholder", None, "descriptive", "secondary-band", "Show SNr->STN directional edge and state-module directional settings."),
            PanelSpec("3j", "SNr->STN PSI beta-low", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/psi/mean-scalar_summary_trans/Beta_low/lmer_model/Phase_raw.pdf", "lmer", "beta-low", "Primary directional extension panel for the state module.", local_filename="3j.pdf"),
            PanelSpec("3k", "State module export", "placeholder", None, "descriptive", "secondary-band", "Show standardized state-wise tables and statistical exports."),
        ],
        "figure4": [
            PanelSpec("4a", "Event-driven overview", "placeholder", None, "descriptive", "secondary-band", "Introduce turn-aligned and cycle-resolved workflows."),
            PanelSpec("4b", "Turn settings and phase definition", "placeholder", None, "descriptive", "secondary-band", "Show Pre, Onset, Offset, and Post windows."),
            PanelSpec("4c", "Representative turn-aligned raw trace", "placeholder", None, "descriptive", "secondary-band", "Illustrate event locking for turning."),
            PanelSpec("4d", "Turn spectral overview", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/periodic/mean-spectral_summary_trans/Phase-Region_raw.pdf", "descriptive", "beta-high", "Turn spectral overview with event-phase context.", local_filename="4d.pdf"),
            PanelSpec("4e", "SNr theta burst duration", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/burst/duration-scalar_summary_trans/Theta/lmer_model/Phase-Region_raw.pdf", "lmer", "secondary-band", "Turn-local extension panel retained despite non-beta band because event alignment is the main question.", local_filename="4e.pdf"),
            PanelSpec("4f", "SNr-STN ciPLV Gamma", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/ciplv/mean-scalar_summary_trans/Gamma/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Turn synchrony extension panel.", local_filename="4f.pdf"),
            PanelSpec("4g", "SNr->STN PSI beta-high", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/psi/mean-scalar_summary_trans/Beta_high/lmer_model/Phase_raw.pdf", "lmer", "beta-high", "Primary beta-high directional panel for turn.", local_filename="4g.pdf"),
            PanelSpec("4h", "Cycle shift schematic", "placeholder", None, "descriptive", "secondary-band", "Show ipsi-to-contra shift and the rationale for interval analysis."),
            PanelSpec("4i", "Cycle local periodic interval", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/periodic/mean-trace_summary_trans_normalized_shift/Theta/Region_norm.pdf", "descriptive", "secondary-band", "Retained periodic cycle interval around 97-4% for SNr theta.", local_filename="4i.pdf", interval_key=("periodic", "Theta", "SNr", "below_0", 97, 4)),
            PanelSpec("4j", "Cycle local aperiodic interval", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/aperiodic/mean-trace_summary_trans_normalized_shift/Offset/Region_norm.pdf", "descriptive", "secondary-band", "Retained aperiodic cycle interval around 3-12% for SNr offset.", local_filename="4j.pdf", interval_key=("aperiodic", "Offset", "SNr", "above_0", 3, 12)),
            PanelSpec("4k", "Cycle coherence interval", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/coherence/mean-trace_summary_trans_normalized_shift/Alpha/Region_norm.pdf", "descriptive", "secondary-band", "Retained synchrony interval for coherence Alpha.", local_filename="4k.pdf", interval_key=("coherence", "Alpha", "SNr-STN", "below_0", 1, 11)),
            PanelSpec("4l", "Cycle wPLI interval", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/wpli/mean-trace_summary_trans_normalized_shift/Alpha/Region_norm.pdf", "descriptive", "secondary-band", "Retained synchrony interval for wPLI Alpha.", local_filename="4l.pdf", interval_key=("wpli", "Alpha", "SNr-STN", "below_0", 99, 11)),
        ],
        "figureS1": [
            PanelSpec("S1a", "Localization atlas overview", "copied", assets["leaddbs_png"], "descriptive", "secondary-band", "Reuse localization asset for the supplementary anatomy atlas.", local_filename="S1a.png"),
            PanelSpec("S1b", "Full contact coverage layout", "placeholder", None, "descriptive", "secondary-band", "Layout for full contact coverage across subjects and hemispheres."),
            PanelSpec("S1c", "Pair taxonomy atlas", "placeholder", None, "descriptive", "secondary-band", "Atlas view of STN-local, SNr-local, synchrony, and directional pair classes."),
        ],
        "figureS2": [
            PanelSpec("S2a", "Rest coherence Theta", "copied", "/Users/mojackhu/Research/pd/summary/table/med/coherence/mean-scalar_summary_trans/Theta/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Rest connectivity atlas entry for coherence Theta.", local_filename="S2a.pdf"),
            PanelSpec("S2b", "Rest wPLI Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/med/wpli/mean-scalar_summary_trans/Beta_low/rlmer_model/Phase_raw.pdf", "descriptive", "beta-low", "Rest connectivity atlas entry for wPLI beta-low.", local_filename="S2b.pdf"),
            PanelSpec("S2c", "Rest ciPLV Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/med/ciplv/mean-scalar_summary_trans/Beta_low/rlmer_model/Phase_raw.pdf", "descriptive", "beta-low", "Rest connectivity atlas entry for ciPLV beta-low.", local_filename="S2c.pdf"),
            PanelSpec("S2d", "Rest PSI Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/med/psi/mean-scalar_summary_trans/Beta_low/lmer_model/Phase_raw.pdf", "lmer", "beta-low", "Rest connectivity atlas entry for PSI beta-low.", local_filename="S2d.pdf"),
            PanelSpec("S2e", "Rest TRGC Delta", "copied", "/Users/mojackhu/Research/pd/summary/table/med/trgc/mean-scalar_summary_trans/Delta/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Rest connectivity atlas entry for TRGC delta.", local_filename="S2e.pdf"),
        ],
        "figureS3": [
            PanelSpec("S3a", "Motor PSI Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/psi/mean-scalar_summary_trans/Beta_low/lmer_model/Phase_raw.pdf", "lmer", "beta-low", "Motor connectivity atlas entry for PSI beta-low.", local_filename="S3a.pdf"),
            PanelSpec("S3b", "Motor PSI Beta_high", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/psi/mean-scalar_summary_trans/Beta_high/rlmer_model/Phase_raw.pdf", "descriptive", "beta-high", "Motor connectivity atlas entry for PSI beta-high.", local_filename="S3b.pdf"),
            PanelSpec("S3c", "Motor ciPLV Alpha", "copied", "/Users/mojackhu/Research/pd/summary/table/motor/ciplv/mean-scalar_summary_trans/Alpha/rlmer_model/Phase_raw.pdf", "descriptive", "secondary-band", "Motor connectivity atlas entry for ciPLV alpha.", local_filename="S3c.pdf"),
            PanelSpec("S3d", "Motor connectivity atlas layout", "placeholder", None, "descriptive", "secondary-band", "Placeholder for combined motor connectivity atlas annotations."),
        ],
        "figureS4": [
            PanelSpec("S4a", "Turn ciPLV Gamma", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/ciplv/mean-scalar_summary_trans/Gamma/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Turn connectivity atlas entry for ciPLV gamma.", local_filename="S4a.pdf"),
            PanelSpec("S4b", "Turn PSI Beta_high", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/psi/mean-scalar_summary_trans/Beta_high/lmer_model/Phase_raw.pdf", "lmer", "beta-high", "Turn connectivity atlas entry for PSI beta-high.", local_filename="S4b.pdf"),
            PanelSpec("S4c", "Turn PLI Gamma", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/pli/mean-scalar_summary_trans/Gamma/lmer_model/Phase_raw.pdf", "lmer", "secondary-band", "Turn connectivity atlas entry for PLI gamma.", local_filename="S4c.pdf"),
            PanelSpec("S4d", "Turn wPLI Gamma", "copied", "/Users/mojackhu/Research/pd/summary/table/turn/wpli/mean-scalar_summary_trans/Gamma/rlmer_model/Phase_raw.pdf", "descriptive", "secondary-band", "Turn connectivity atlas entry for wPLI gamma.", local_filename="S4d.pdf"),
        ],
        "figureS5": [
            PanelSpec("S5a", "Cycle periodic Theta trace", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/periodic/mean-trace_summary_trans_normalized_shift/Theta/Region_norm.pdf", "descriptive", "secondary-band", "Cycle local atlas entry for periodic theta.", local_filename="S5a.pdf", interval_key=("periodic", "Theta", "SNr", "below_0", 97, 4)),
            PanelSpec("S5b", "Cycle periodic Alpha trace", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/periodic/mean-trace_summary_trans_normalized_shift/Alpha/Region_norm.pdf", "descriptive", "secondary-band", "Cycle local atlas entry for periodic alpha.", local_filename="S5b.pdf", interval_key=("periodic", "Alpha", "SNr", "below_0", 1, 9)),
            PanelSpec("S5c", "Cycle aperiodic Offset trace", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/aperiodic/mean-trace_summary_trans_normalized_shift/Offset/Region_norm.pdf", "descriptive", "secondary-band", "Cycle local atlas entry for aperiodic offset.", local_filename="S5c.pdf", interval_key=("aperiodic", "Offset", "SNr", "above_0", 3, 12)),
            PanelSpec("S5d", "Cycle aperiodic Exponent trace", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/aperiodic/mean-trace_summary_trans_normalized_shift/Exponent/Region_norm.pdf", "descriptive", "secondary-band", "Cycle local atlas entry for aperiodic exponent.", local_filename="S5d.pdf"),
            PanelSpec("S5e", "Cycle raw-power atlas layout", "placeholder", None, "descriptive", "secondary-band", "Placeholder for optional raw-power local atlas annotation."),
        ],
        "figureS6": [
            PanelSpec("S6a", "Cycle coherence Alpha", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/coherence/mean-trace_summary_trans_normalized_shift/Alpha/Region_norm.pdf", "descriptive", "secondary-band", "Cycle connectivity atlas entry for coherence alpha.", local_filename="S6a.pdf", interval_key=("coherence", "Alpha", "SNr-STN", "below_0", 1, 11)),
            PanelSpec("S6b", "Cycle coherence Theta", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/coherence/mean-trace_summary_trans_normalized_shift/Theta/Region_norm.pdf", "descriptive", "secondary-band", "Cycle connectivity atlas entry for coherence theta.", local_filename="S6b.pdf", interval_key=("coherence", "Theta", "SNr-STN", "below_0", 57, 65)),
            PanelSpec("S6c", "Cycle wPLI Alpha", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/wpli/mean-trace_summary_trans_normalized_shift/Alpha/Region_norm.pdf", "descriptive", "secondary-band", "Cycle connectivity atlas entry for wPLI alpha.", local_filename="S6c.pdf", interval_key=("wpli", "Alpha", "SNr-STN", "below_0", 99, 11)),
            PanelSpec("S6d", "Cycle wPLI Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/wpli/mean-trace_summary_trans_normalized_shift/Beta_low/Region_norm.pdf", "descriptive", "beta-low", "Cycle connectivity atlas entry for wPLI beta-low.", local_filename="S6d.pdf", interval_key=("wpli", "Beta_low", "SNr-STN", "below_0", 89, 98)),
            PanelSpec("S6e", "Cycle PLV Beta_low", "copied", "/Users/mojackhu/Research/pd/summary/table/cycle/plv/mean-trace_summary_trans_normalized_shift/Beta_low/Region_norm.pdf", "descriptive", "beta-low", "Cycle connectivity atlas entry for PLV beta-low.", local_filename="S6e.pdf"),
            PanelSpec("S6f", "Cycle ciPLV/TRGC support layout", "placeholder", None, "descriptive", "secondary-band", "Placeholder describing toolbox-supported ciPLV and TRGC families even when retained intervals are absent."),
        ],
        "figureS7": [
            PanelSpec("S7a", "App configuration overview", "placeholder", None, "descriptive", "secondary-band", "Show module configuration screens and required settings."),
            PanelSpec("S7b", "Workflow schematic reuse", "copied", assets["workflow_spec_md"], "descriptive", "secondary-band", "Reuse the workflow specification as a supplementary app/export atlas panel.", local_filename="S7b.md"),
            PanelSpec("S7c", "Export artifact layout", "placeholder", None, "descriptive", "secondary-band", "Show CSV/XLSX exports, stats folders, and figure-manifest outputs."),
        ],
    }


def table_header_comments() -> dict[str, dict[str, str]]:
    return {
        "Table1": {
            "Age": "Age is not present in the current PD workspace. Populate from the final clinical spreadsheet before submission.",
            "Sex": "Sex is not present in the current PD workspace. Populate from the final clinical spreadsheet before submission.",
            "DiseaseDurationYears": "Disease duration is not present in the current PD workspace. Populate from the final clinical spreadsheet before submission.",
            "LEDD": "LEDD is not present in the current PD workspace. Populate from the final clinical spreadsheet before submission.",
            "UPDRSIII_Off": "Confirm preoperative or medication-OFF UPDRS-III source for the final submission table.",
            "UPDRSIII_On": "Confirm medication-ON UPDRS-III source for the final submission table.",
        },
        "TableS1": {
            "Age": "Populate full clinical metadata from the locked clinical source workbook.",
            "Sex": "Populate full clinical metadata from the locked clinical source workbook.",
            "DiseaseDurationYears": "Populate full clinical metadata from the locked clinical source workbook.",
            "LEDD": "Populate full clinical metadata from the locked clinical source workbook.",
            "HoehnYahrOff": "Populate full clinical metadata from the locked clinical source workbook.",
            "FOGPhenotype": "Confirm gait phenotype / FOG classification source before submission.",
        },
    }


def _load_paradigm_table(project_root: Path) -> pd.DataFrame:
    workbook = summary_root(project_root) / "cohort" / "subj_paradigm.xlsx"
    raw = pd.read_excel(workbook)
    return normalize_subj_paradigm_frame(raw)


def _load_channel_coords(project_root: Path) -> pd.DataFrame:
    coords = summary_root(project_root) / "cohort" / "channel_coords.csv"
    frame = pd.read_csv(coords)
    frame["SubjectCode"] = frame["Subject"].str.lower()
    frame["Side"] = frame["Channel"].astype(str).str.startswith(("0", "1", "2")).map(
        lambda is_left: "L" if is_left else "R"
    )
    return frame


def build_table1(project_root: Path) -> pd.DataFrame:
    paradigm = _load_paradigm_table(project_root)
    coords = _load_channel_coords(project_root)
    counts = (
        coords.groupby(["SubjectCode", "Region"]).size().unstack(fill_value=0).reset_index()
    )
    table = paradigm.merge(counts, on="SubjectCode", how="left")
    for col in ["Age", "Sex", "DiseaseDurationYears", "LEDD", "UPDRSIII_Off", "UPDRSIII_On"]:
        table[col] = ""
    for region in ["SNr", "STN", "Mid"]:
        if region not in table.columns:
            table[region] = 0
    table["SNrChannels"] = table["SNr"].fillna(0).astype(int)
    table["STNChannels"] = table["STN"].fillna(0).astype(int)
    table["MidChannels"] = table["Mid"].fillna(0).astype(int)
    table["Notes"] = "Populate missing clinical columns from the final cohort workbook."
    ordered = [
        "SubjectIndex",
        "SubjectCode",
        "SubjectName",
        "Age",
        "Sex",
        "DiseaseDurationYears",
        "LEDD",
        "UPDRSIII_Off",
        "UPDRSIII_On",
        "Sit",
        "Stand",
        "Walk",
        "Turn",
        "Cycle",
        "FoG",
        "Med",
        "Pain",
        "SNrChannels",
        "STNChannels",
        "MidChannels",
        "Notes",
    ]
    return table.loc[:, ordered]


def build_table2(project_root: Path) -> pd.DataFrame:
    _ = project_root
    return pd.DataFrame(
        [
            {
                "PairClass": "STN local pair",
                "RegionScope": "STN",
                "SelectionRule": "Adjacency-based bipolar pair whose localization falls inside STN.",
                "UsedInRest": "Yes",
                "UsedInMotor": "Yes",
                "UsedInTurn": "Yes",
                "UsedInCycle": "Yes",
                "FeatureFamilies": "periodic; aperiodic; burst",
                "OutputType": "Local single-channel summaries",
            },
            {
                "PairClass": "SNr local pair",
                "RegionScope": "SNr",
                "SelectionRule": "Adjacency-based bipolar pair whose localization falls inside SNr.",
                "UsedInRest": "Yes",
                "UsedInMotor": "Yes",
                "UsedInTurn": "Yes",
                "UsedInCycle": "Yes",
                "FeatureFamilies": "periodic; aperiodic; burst",
                "OutputType": "Local single-channel summaries",
            },
            {
                "PairClass": "SNr-STN synchrony edge",
                "RegionScope": "SNr-STN",
                "SelectionRule": "Cross-region edge linking one SNr pair and one STN pair for undirected connectivity.",
                "UsedInRest": "Yes",
                "UsedInMotor": "Yes",
                "UsedInTurn": "Yes",
                "UsedInCycle": "Yes",
                "FeatureFamilies": "coherence; wPLI; ciPLV; PLI; PLV",
                "OutputType": "Undirected network summaries",
            },
            {
                "PairClass": "SNr->STN directional edge",
                "RegionScope": "SNr->STN",
                "SelectionRule": "Ordered cross-region edge linking one SNr pair to one STN pair for directional connectivity.",
                "UsedInRest": "Yes",
                "UsedInMotor": "Yes",
                "UsedInTurn": "Yes",
                "UsedInCycle": "Yes",
                "FeatureFamilies": "TRGC; PSI",
                "OutputType": "Directed network summaries",
            },
        ]
    )


def build_table3(project_root: Path) -> pd.DataFrame:
    _ = project_root
    return pd.DataFrame(
        [
            {
                "Module": "Rest module",
                "ParadigmType": "Medication contrast",
                "AlignmentType": "Condition contrast (Off vs On)",
                "LocalFeatures": "periodic; aperiodic; burst",
                "SynchronyFeatures": "coherence; wPLI; ciPLV",
                "DirectionalFeatures": "TRGC; PSI",
                "MainTextAnchors": "STN beta-low periodic; STN beta-low burst; SNr-STN coherence Theta; SNr->STN TRGC Delta",
                "SupplementaryOutputs": "Full connectivity atlas and non-retained frequency bands",
                "ExportArtifact": "summary/table/med + figure2/ manifest",
            },
            {
                "Module": "State module",
                "ParadigmType": "Standing-to-walking state contrast",
                "AlignmentType": "Block/state alignment",
                "LocalFeatures": "periodic; aperiodic; burst",
                "SynchronyFeatures": "coherence; wPLI; ciPLV",
                "DirectionalFeatures": "TRGC; PSI",
                "MainTextAnchors": "STN beta-low burst rate; STN beta-low burst occupation; SNr->STN PSI beta-low",
                "SupplementaryOutputs": "Motor connectivity atlas and secondary-band context",
                "ExportArtifact": "summary/table/motor + figure3/ manifest",
            },
            {
                "Module": "Turn event module",
                "ParadigmType": "Discrete event-driven analysis",
                "AlignmentType": "Turn-phase alignment",
                "LocalFeatures": "periodic; aperiodic; burst",
                "SynchronyFeatures": "coherence; wPLI; ciPLV",
                "DirectionalFeatures": "TRGC; PSI",
                "MainTextAnchors": "SNr theta burst duration; SNr-STN ciPLV Gamma; SNr->STN PSI beta-high",
                "SupplementaryOutputs": "Turn connectivity atlas across all supported families",
                "ExportArtifact": "summary/table/turn + figure4/ turn panels",
            },
            {
                "Module": "Cycle interval module",
                "ParadigmType": "Continuous phase-resolved analysis",
                "AlignmentType": "Circular gait-cycle alignment after ipsi-to-contra shift",
                "LocalFeatures": "periodic; aperiodic; raw_power",
                "SynchronyFeatures": "coherence; wPLI; ciPLV; PLI; PLV",
                "DirectionalFeatures": "TRGC; PSI",
                "MainTextAnchors": "SNr periodic Theta; SNr aperiodic Offset; SNr-STN coherence Alpha; SNr-STN wPLI Alpha",
                "SupplementaryOutputs": "Full cycle local and connectivity interval atlases",
                "ExportArtifact": "summary/table/cycle/interval + figure4/ cycle panels",
            },
        ]
    )


def build_table_s1(project_root: Path) -> pd.DataFrame:
    base = build_table1(project_root).copy()
    extra_cols = ["HoehnYahrOff", "FOGPhenotype", "TremorDominant", "ClinicalNotes"]
    for col in extra_cols:
        base[col] = ""
    base["ClinicalNotes"] = "Populate from final locked clinical spreadsheet."
    return base


def build_table_s2(project_root: Path) -> pd.DataFrame:
    paradigm = _load_paradigm_table(project_root)
    coords = _load_channel_coords(project_root)
    rows: list[dict[str, Any]] = []
    for module, column in [
        ("Rest / medication", "Med"),
        ("Standing-to-walking", "Walk"),
        ("Turn", "Turn"),
        ("Cycle", "Cycle"),
        ("Pain", "Pain"),
    ]:
        available = int((paradigm[column] == "Yes").sum()) if column in paradigm else 0
        rows.append(
            {
                "Module": module,
                "AvailabilityColumn": column,
                "SubjectsAvailable": available,
                "SubjectsWithSNrChannels": int(coords.loc[coords["Region"] == "SNr", "SubjectCode"].nunique()),
                "SubjectsWithSTNChannels": int(coords.loc[coords["Region"] == "STN", "SubjectCode"].nunique()),
                "TotalChannelRows": int(len(coords)),
                "Notes": "Feature-specific analyzable counts vary by metric family and are summarized in Table S4 and Table S5.",
            }
        )
    return pd.DataFrame(rows)


def build_table_s3(project_root: Path) -> pd.DataFrame:
    paradigm = _load_paradigm_table(project_root).loc[:, ["SubjectCode", "SubjectName"]]
    coords = _load_channel_coords(project_root).merge(paradigm, on="SubjectCode", how="left")
    ordered = [
        "Subject",
        "SubjectCode",
        "SubjectName",
        "Channel",
        "Side",
        "Region",
        "MNI_x",
        "MNI_y",
        "MNI_z",
    ]
    coords["Subject"] = coords["SubjectCode"]
    return coords.loc[:, ordered].sort_values(["SubjectCode", "Channel"]).reset_index(drop=True)


def _summarize_tukey_group(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str, str, str], dict[str, Any]] = {}
    for rec in records:
        key = (
            rec["Section"],
            rec["MetricPath"],
            rec["Band"],
            rec["WorkbookStem"],
        )
        row = grouped.setdefault(
            key,
            {
                "Section": rec["Section"],
                "MetricPath": rec["MetricPath"],
                "Band": rec["Band"],
                "WorkbookStem": rec["WorkbookStem"],
                "lmer_min_p": None,
                "lmer_sig_contrasts": 0,
                "lmer_pdf": "",
                "rlmer_min_p": None,
                "rlmer_sig_contrasts": 0,
                "rlmer_pdf": "",
            },
        )
        model_prefix = "lmer" if rec["Model"] == "lmer_model" else "rlmer"
        row[f"{model_prefix}_min_p"] = rec["MinP"]
        row[f"{model_prefix}_sig_contrasts"] = rec["SigContrasts"]
        row[f"{model_prefix}_pdf"] = rec["RawPdf"]
    out: list[dict[str, Any]] = []
    for row in grouped.values():
        row["dual_model_support"] = (
            "yes"
            if (row["lmer_sig_contrasts"] > 0 and row["rlmer_sig_contrasts"] > 0)
            else "no"
        )
        out.append(row)
    return sorted(out, key=lambda r: (r["Section"], r["MetricPath"], r["Band"], r["WorkbookStem"]))


def build_table_s4(project_root: Path, chosen_pdfs: set[str]) -> pd.DataFrame:
    records: list[dict[str, Any]] = []
    root = summary_table_root(project_root)
    for csv_path in root.rglob("*_tukey.csv"):
        rel = csv_path.relative_to(root)
        if not rel.parts or rel.parts[0] not in {"med", "motor", "turn"}:
            continue
        parts = rel.parts
        model_idx = None
        for idx, part in enumerate(parts):
            if part in {"lmer_model", "rlmer_model"}:
                model_idx = idx
                break
        if model_idx is None or model_idx < 2:
            continue
        frame = pd.read_csv(csv_path)
        if "p_tukey" not in frame.columns:
            continue
        model = parts[model_idx]
        band = parts[model_idx - 1]
        metric_path = "/".join(parts[1 : model_idx - 1])
        raw_pdf = csv_path.with_name(csv_path.name.replace("_tukey.csv", "_raw.pdf"))
        records.append(
            {
                "Section": parts[0],
                "MetricPath": metric_path,
                "Band": band,
                "WorkbookStem": csv_path.name.replace("_tukey.csv", ""),
                "Model": model,
                "MinP": float(frame["p_tukey"].min()) if len(frame) else None,
                "SigContrasts": int((frame["p_tukey"] < 0.05).sum()),
                "RawPdf": raw_pdf.as_posix(),
            }
        )
    summary = pd.DataFrame(_summarize_tukey_group(records))
    if summary.empty:
        return summary
    summary["selected_for_main_text"] = summary["lmer_pdf"].isin(chosen_pdfs).map(
        lambda value: "yes" if value else "no"
    )
    return summary


def build_table_s5(project_root: Path) -> pd.DataFrame:
    path = summary_table_root(project_root) / "cycle" / "interval" / "cycle_timepoint_intervals.csv"
    return pd.read_csv(path).sort_values(
        ["Metric", "Band", "Region", "direction", "start_pct", "end_pct"]
    ).reset_index(drop=True)


def build_table_s6(project_root: Path) -> pd.DataFrame:
    _ = project_root
    return pd.DataFrame(
        [
            {
                "FeatureFamily": "periodic",
                "FeatureType": "Local",
                "MainTextPlacement": "Figures 2-4",
                "SupplementPlacement": "S5",
                "NarrativeRole": "Primary beta-centered oscillatory narrative and event-driven extensions",
            },
            {
                "FeatureFamily": "aperiodic",
                "FeatureType": "Local",
                "MainTextPlacement": "Figure 4",
                "SupplementPlacement": "S5",
                "NarrativeRole": "Cycle-resolved local extension beyond canonical beta markers",
            },
            {
                "FeatureFamily": "burst",
                "FeatureType": "Local",
                "MainTextPlacement": "Figures 2-4",
                "SupplementPlacement": "S2-S4",
                "NarrativeRole": "Primary dynamic validation family for rest, motor, and turn modules",
            },
            {
                "FeatureFamily": "coherence",
                "FeatureType": "Synchrony",
                "MainTextPlacement": "Figures 2 and 4",
                "SupplementPlacement": "S2-S6",
                "NarrativeRole": "Primary undirected network family for rest and cycle extensions",
            },
            {
                "FeatureFamily": "wPLI",
                "FeatureType": "Synchrony",
                "MainTextPlacement": "Figure 4",
                "SupplementPlacement": "S2-S6",
                "NarrativeRole": "Lag-sensitive synchrony family retained for cycle interval discovery",
            },
            {
                "FeatureFamily": "ciPLV",
                "FeatureType": "Synchrony",
                "MainTextPlacement": "Figure 4",
                "SupplementPlacement": "S2-S6",
                "NarrativeRole": "Turn synchrony extension and toolbox-supported connectivity family",
            },
            {
                "FeatureFamily": "TRGC",
                "FeatureType": "Directional",
                "MainTextPlacement": "Figure 2",
                "SupplementPlacement": "S2-S6",
                "NarrativeRole": "Directional family highlighted in rest and carried as supported capability elsewhere",
            },
            {
                "FeatureFamily": "PSI",
                "FeatureType": "Directional",
                "MainTextPlacement": "Figures 3 and 4",
                "SupplementPlacement": "S2-S6",
                "NarrativeRole": "Directional family retained for motor and turn module anchors",
            },
        ]
    )


def _write_excel_table(
    frame: pd.DataFrame,
    path: Path,
    *,
    sheet_name: str,
    header_comments: dict[str, str] | None = None,
) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        frame.to_excel(writer, index=False, sheet_name=sheet_name)
        ws = writer.book[sheet_name]
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        for idx, column in enumerate(frame.columns, start=1):
            width = max(len(str(column)) + 2, 14)
            ws.column_dimensions[get_column_letter(idx)].width = min(width, 28)
            if header_comments and column in header_comments:
                ws.cell(row=1, column=idx).comment = Comment(header_comments[column], COMMENT_AUTHOR)


def _panel_filename(panel: PanelSpec) -> str:
    if panel.local_filename is not None:
        return panel.local_filename
    suffix = ".md" if panel.status == "placeholder" else Path(panel.source_path or "").suffix
    return f"{panel.panel_id}{suffix}"


def _placeholder_panel_text(panel: PanelSpec) -> str:
    return (
        f"# Panel {panel.panel_id}\n\n"
        f"## Role\n\n{panel.panel_role}\n\n"
        f"## Layout notes\n\n{panel.notes}\n"
    )


def _figure_layout_text(figure_name: str, panels: list[PanelSpec]) -> str:
    lines = [f"# {figure_name.capitalize()} Layout", "", "## Panel order", ""]
    for panel in panels:
        lines.append(f"- `{panel.panel_id}`: {panel.panel_role}")
        lines.append(f"  Notes: {panel.notes}")
    return "\n".join(lines) + "\n"


def build_figures(project_root: Path) -> dict[str, Any]:
    specs = figure_specs(project_root)
    figure_root = _figure_table_root(project_root)
    interval_catalog = load_cycle_interval_catalog(project_root)
    manifest: dict[str, Any] = {}
    for figure_name, panels in specs.items():
        out_dir = figure_root / figure_name
        out_dir.mkdir(parents=True, exist_ok=True)
        manifest_rows: list[dict[str, Any]] = []
        for panel in panels:
            filename = _panel_filename(panel)
            dest = out_dir / filename
            support = panel.dual_model_support
            support_note = panel.dual_model_notes
            if panel.status == "copied" and panel.source_path:
                src = Path(panel.source_path)
                if src.suffix.lower() in {".pdf", ".png", ".tif", ".ai", ".md"} and src.exists():
                    if src.suffix.lower() == ".md" and dest.suffix.lower() == ".md":
                        dest.write_text(src.read_text(encoding="utf-8"), encoding="utf-8")
                    else:
                        shutil.copy2(src, dest)
                else:
                    dest.write_text(_placeholder_panel_text(panel), encoding="utf-8")
                if panel.interval_key is not None:
                    support, support_note = evaluate_interval_dual_support(interval_catalog, panel.interval_key)
                elif panel.display_model == "lmer":
                    support, support_note = evaluate_scalar_dual_support(panel.source_path)
            else:
                dest.write_text(_placeholder_panel_text(panel), encoding="utf-8")
            manifest_rows.append(
                {
                    "panel_id": panel.panel_id,
                    "status": panel.status,
                    "panel_role": panel.panel_role,
                    "source_path": panel.source_path,
                    "local_filename": filename,
                    "display_model": panel.display_model,
                    "dual_model_support": support,
                    "band_priority": panel.band_priority,
                    "notes": panel.notes,
                    "support_notes": support_note,
                }
            )
        (out_dir / "layout.md").write_text(
            _figure_layout_text(figure_name, panels), encoding="utf-8"
        )
        (out_dir / f"{figure_name}.yml").write_text(
            yaml.safe_dump(manifest_rows, sort_keys=False, allow_unicode=False),
            encoding="utf-8",
        )
        manifest[figure_name] = manifest_rows
    return manifest


def _build_reference_list() -> list[str]:
    return [
        "1. Kuhn AA, Williams D, Kupsch A, Limousin P, Hariz M, Schneider GH, et al. Event-related beta desynchronization in human subthalamic nucleus correlates with motor performance. Brain. 2004;127:735-746.",
        "2. Torrecillos F, Tinkhauser G, Fischer P, Green AL, Aziz TZ, Foltynie T, et al. Modulation of beta bursts in the subthalamic nucleus predicts motor performance. J Neurosci. 2018;38:8905-8917.",
        "3. Lofredi R, Neumann WJ, Bock A, Horn A, Huebl J, Siegert S, et al. Dopamine-dependent scaling of subthalamic gamma bursts with movement velocity in patients with Parkinson's disease. Elife. 2018;7:e31895.",
        "4. Tinkhauser G, Pogosyan A, Tan H, Herz DM, Kuhn AA, Brown P. The modulatory effect of adaptive deep brain stimulation on beta bursts in Parkinson's disease. Brain. 2017;140:1053-1067.",
        "5. Neudorfer C, et al. Lead-DBS v3.0: Mapping deep brain stimulation effects to local anatomy and global networks. Neuroimage. 2023;268:119862.",
        "6. Oxenford S, et al. WarpDrive: Improving spatial normalization using manual refinements. Med Image Anal. 2024;91:103041.",
        "7. Donoghue T, Haller M, Peterson EJ, Varma P, Sebastian P, Gao R, et al. Parameterizing neural power spectra into periodic and aperiodic components. Nat Neurosci. 2020;23:1655-1665.",
        "8. Vinck M, Oostenveld R, van Wingerden M, Battaglia F, Pennartz CMA. An improved index of phase-synchronization for electrophysiological data in the presence of volume-conduction, noise and sample-size bias. Neuroimage. 2011;55:1548-1565.",
        "9. Bruna R, Maestu F, Pereda E. Phase locking value revisited: teaching new tricks to an old dog. J Neural Eng. 2018;15:056011.",
        "10. Barnett L, Seth AK. Granger causality for state-space models. Phys Rev E Stat Nonlin Soft Matter Phys. 2015;91:040101.",
        "11. Nolte G, Ziehe A, Nikulin VV, Schlögl A, Krämer N, Brismar T, et al. Robustly estimating the flow direction of information in complex physical systems. Phys Rev Lett. 2008;100:234101.",
        "12. Postuma RB, Berg D, Stern M, Poewe W, Olanow CW, Oertel W, et al. MDS clinical diagnostic criteria for Parkinson's disease. Mov Disord. 2015;30:1591-1601.",
        "13. Defer GL, Widner H, Marie RM, Remy P, Levivier M. Core assessment program for surgical interventional therapies in Parkinson's disease (CAPSIT-PD). Mov Disord. 1999;14:572-584.",
        "14. Hoehn MM, Yahr MD. Parkinsonism: onset, progression and mortality. Neurology. 1967;17:427-442.",
        "15. Ewert S, Plettig P, Li N, Chakravarty MM, Collins DL, Herrington TM, et al. Toward defining deep brain stimulation targets in MNI space: A subcortical atlas based on multimodal MRI, histology and structural connectivity. Neuroimage. 2018;170:271-282.",
        "16. Husch A, Petersen MV, Gemmar P, Goncalves J, Hertel F. PaCER - A fully automated method for electrode trajectory and contact reconstruction in deep brain stimulation. Neuroimage Clin. 2018;17:80-89.",
        "17. Jas M, Engemann DA, Bekhti Y, Raimondo F, Gramfort A. Autoreject: Automated artifact rejection for MEG and EEG data. Neuroimage. 2017;159:417-429.",
        "18. Gerster M, Waterstraat G, Litvak V, Lehnertz K, Schnitzler A, Florin E, et al. Separating neural oscillations from aperiodic 1/f activity: challenges and recommendations. Neuroinformatics. 2022;20:991-1012.",
        "19. Tan H, Fischer P, Shah SA, Vidaurre D, Woolrich MW, Brown P. Decoding movement states in stepping cycles based on subthalamic LFPs in Parkinsonian patients. Annu Int Conf IEEE Eng Med Biol Soc. 2018;2018:1384-1387.",
        "20. Georgiades MJ, Shine JM, Gilat M, McMaster J, Owler B, Mahant N, et al. Subthalamic nucleus activity during cognitive load and gait dysfunction in Parkinson's disease. Mov Disord. 2023;38:1549-1554.",
        "21. Shi L, et al. Comparison of oscillatory activity in substantia nigra pars reticulata between Parkinson's disease and dystonia. NPJ Parkinsons Dis. 2025;11:109.",
    ]


def _highlights_text() -> str:
    return """# Highlights

- A single anatomy-guided toolbox supported medication, locomotor-state, turn-event, and gait-cycle analyses in one STN+SNr cohort.
- The rest module reproduced canonical STN beta-low physiology while extending the same workflow to SNr-STN coherence and SNr->STN directionality.
- The state module captured locomotor-state effects through beta-centered burst features and nigro-subthalamic directional coupling.
- The event-driven modules combined discrete turn alignment with continuous cycle-resolved interval discovery after ipsi-to-contra phase shifting.
- Unified exports linked local, synchrony, and directional features to manuscript-ready figure and table assets for Brain Stimulation submission.
"""


def _title_page_text() -> str:
    return """# Title Page

## Title

An anatomy-guided toolbox for multi-paradigm STN-SNr electrophysiological analysis in Parkinson's disease

## Running title

Toolbox for STN-SNr multi-paradigm electrophysiology

## Authors

[[COMMENT: Populate the final author list, affiliations, and equal-contribution notes from the locked contributor spreadsheet.]]

## Corresponding author

[[COMMENT: Add corresponding author name, postal address, email, and phone number.]]

## Article type

Original Article

## Keywords

Parkinson's disease; deep brain stimulation; substantia nigra pars reticulata; subthalamic nucleus; beta oscillations; toolbox

## Word-count and figure-count checkpoints

[[COMMENT: Confirm Brain Stimulation word-count, abstract-count, and artwork-count limits before submission freeze.]]
"""


def _main_manuscript_text() -> str:
    refs = _build_reference_list()
    references_block = "\n".join(refs)
    return f"""# Main Manuscript

## Structured Abstract

### Objective

Deep brain stimulation sensing studies in Parkinson's disease often rely on separate scripts for preprocessing, alignment, local spectral analysis, synchrony analysis, and directional modeling. We developed an anatomy-guided toolbox to analyze subthalamic nucleus (STN) and substantia nigra pars reticulata (SNr) electrophysiology across multiple paradigms within one standardized workflow, with beta-low and beta-high features prioritized as the primary electrophysiological language.

### Methods

The toolbox was applied to one chronic STN+SNr sensing cohort and organized around four analysis modes: medication contrast, locomotor-state contrast, discrete turn-event alignment, and continuous gait-cycle interval analysis. Anatomy-guided STN-local, SNr-local, synchrony, and directional channel definitions were linked to a unified preprocessing and export pipeline. Feature families included periodic and aperiodic components, burst dynamics, coherence, weighted phase-lag index (wPLI), corrected imaginary phase-locking value (ciPLV), time-reversed Granger causality (TRGC), and phase slope index (PSI). Group-level scalar analyses followed the current paper workflow with linear mixed-effects modeling and robust linear mixed-effects confirmation, while continuous gait-cycle analyses used pointwise mixed models and retained contiguous cycle intervals [1-11].

### Results

The rest module reproduced canonical STN beta-low periodic and burst effects while extending the same workflow to resting SNr-STN synchrony and directional features. The state module captured standing-to-walking modulation primarily through beta-low burst summaries and nigro-subthalamic directional coupling. The event-driven modules demonstrated two complementary capabilities: turn-aligned analysis revealed SNr local and network changes, including beta-high directional coupling, whereas cycle-resolved analysis identified sustained SNr and SNr-STN intervals after ipsi-to-contra phase alignment, even when retained intervals extended beyond the beta bands. Across paradigms, the toolbox generated analysis-ready scalar outputs, interval summaries, and figure-ready panel assets from one standardized output grammar.

### Conclusions

This toolbox supports anatomy-guided, multi-paradigm, multi-feature electrophysiological analysis within one exportable workflow. By combining beta-centered validation with event-driven extensions to SNr and SNr-STN features, it provides a practical submission-ready framework for translational DBS sensing studies.

## Introduction

Parkinson's disease electrophysiology studies have repeatedly shown that beta oscillations and beta bursts in the STN track medication state and motor impairment [1-4]. However, translational sensing studies increasingly require analyses that extend beyond a single nucleus, a single feature family, or a single behavioral alignment scheme. In practice, many projects still bridge that gap by chaining together separate preprocessing scripts, anatomy utilities, connectivity code, and figure-specific notebooks. That fragmentation makes it difficult to compare results across paradigms and harder still to export reproducible, manuscript-ready outputs.

The SNr is a compelling extension target for this type of toolbox work because it sits within the same therapeutic circuit while carrying complementary motor and gait information [5,21]. Combining STN and SNr recordings also creates a natural need for anatomy-guided local pairs, undirected SNr-STN synchrony edges, and directed SNr->STN edges. Those requirements argue for a workflow in which localization, preprocessing, alignment, feature extraction, statistics, and export are designed together rather than retrofitted after feature computation.

We therefore organized the present manuscript around toolbox function rather than around any single biomarker. The resulting workflow supports four analysis modes within one cohort: medication contrast, locomotor-state contrast, discrete turn-event alignment, and continuous circular gait-cycle alignment. Across those paradigms, the toolbox supports local periodic and aperiodic summaries, burst features, synchrony metrics such as coherence, wPLI, and ciPLV, and directional measures including TRGC and PSI [7-11]. The core design principle is that beta-low and beta-high remain the primary physiological language, while other bands are retained as secondary context or event-driven extensions.

The manuscript is therefore structured as a validation-and-extension study of the toolbox itself. Resting analyses first test whether the workflow reproduces canonical STN beta-low physiology. State analyses then ask whether the same workflow captures locomotor-state effects using dynamic local and directional features. Finally, event-driven analyses test whether the toolbox can move beyond block contrasts to reveal discrete turn-locked and continuous gait-cycle structure that fixed-window approaches might miss.

## Results

### Cohort and toolbox architecture

The manuscript is built around a single anatomy-guided STN+SNr sensing workflow applied to one chronic Parkinson's disease cohort. The workflow begins with standardized recording import and region-aware localization, then proceeds through preprocessing, paradigm-specific alignment, feature extraction, group-level statistics, and export of manuscript-ready tables and figure panels (Fig. 1). Four alignment modes are supported within the same software logic: condition contrast, state/block alignment, discrete event alignment, and continuous circular phase alignment. [[COMMENT: Confirm the final cohort descriptor sentence, including the exact analyzable patient count used in the main manuscript title page and Table 1.]]

### Rest module: canonical validation plus resting nigral and nigro-subthalamic extension

Using the rest module, the toolbox reproduced the canonical medication-sensitive STN beta-low signature expected from Parkinson's disease sensing studies (Fig. 2). In the main-text periodic and burst panels, beta-low periodic power, beta-low burst rate, and beta-low burst occupation provided the clearest validation anchors, consistent with prior STN literature [1-4]. These beta-low effects were prioritized in both the narrative and the figure sequence to make the module's validation logic explicit.

The same rest workflow also extended naturally to nigral and nigro-subthalamic readouts. Secondary-band panels highlighted medication-sensitive synchrony and directional effects in the same anatomy-guided network, demonstrating that the rest module was not restricted to local STN physiology. In the present main-text selection, SNr-STN coherence in the theta range and SNr->STN directional coupling in the delta range served as retained network extensions because both the standard and robust mixed-effects workflows supported those outputs. This organization allowed the figure to answer a toolbox question rather than a single-feature question: can one rest module reproduce canonical STN beta-low physiology and also expose network structure without changing analysis logic?

### State module: locomotor-state validation plus nigro-subthalamic directional coupling

The state module was designed to test whether the toolbox could move beyond resting contrasts and resolve locomotor-state effects with the same export grammar. Across standing and walking, dynamic beta-low burst summaries outperformed a simple average-power framing as the primary local readout (Fig. 3). This made beta-low burst rate and beta-low burst occupation the natural main-text anchors for the state module, again keeping the manuscript's electrophysiological narrative centered on beta features.

The state module also provided a retained directional extension through SNr->STN PSI in the beta-low range. This result was particularly important for the toolbox story because it showed that the same cohort and the same anatomical labeling framework could support both local dynamic summaries and directional network outputs under a state-aware alignment strategy. In other words, the module did not merely reproduce a known standing-versus-walking effect; it demonstrated that a unified toolbox can move from classic beta validation to nigro-subthalamic directional analysis without switching pipelines.

### Event-driven modules: turn-aligned and cycle-resolved discovery

The event-driven modules provided the clearest differentiator for the toolbox. The turn module addressed discrete event alignment, whereas the cycle module addressed continuous circular alignment after ipsi-to-contra shifting. Together, these analyses asked whether the toolbox could reveal SNr and SNr-STN structure that would be missed by conventional fixed-window approaches.

In the turn analysis, beta-high directionality served as the primary electrophysiological anchor for the module. SNr->STN PSI in the beta-high range was retained as the main directional panel, while SNr theta burst duration and SNr-STN ciPLV gamma acted as local and synchrony extensions, respectively (Fig. 4). This combination kept the turn story grounded in beta-centered interpretation while still showing that event-aligned SNr and SNr-STN features are accessible within the same workflow.

The cycle module addressed a different problem. Instead of forcing all effects into fixed heel-strike or toe-off windows, the toolbox shifted ipsilateral traces into the contralateral frame, smoothed and downsampled the cycle axis, then detected contiguous intervals that significantly deviated from the normalized cycle baseline. In the current dataset, retained cycle intervals were strongest in secondary bands rather than in beta-low or beta-high, with local periodic theta, local aperiodic offset, coherence alpha, and wPLI alpha emerging as the main interval examples (Fig. 4). We therefore frame these cycle findings as event-driven extensions discovered within a beta-centered toolbox, rather than as replacements for the canonical beta narrative established by the rest and state modules.

### Unified export and toolbox utility

Across all modules, the same anatomy-guided labels, model-selection logic, and export grammar were preserved. The workflow generated summary tables, modeled scalar panels, interval catalogues, and figure manifests in one submission-oriented structure rather than in figure-specific script branches. That portability across paradigms, together with the ability to bridge local, synchrony, and directional families, is the central practical contribution of the toolbox.

## Discussion

The present manuscript was intentionally organized as a toolbox paper rather than as a single-feature physiology paper. That distinction matters because the main contribution is not simply that STN beta-low changes with medication or that SNr-STN network features appear during gait. Instead, the central result is that one anatomy-guided workflow can carry a study from canonical validation through state-aware and event-driven analyses while preserving a shared set of outputs and statistical rules.

This framing also clarifies why beta-low and beta-high remain the primary electrophysiological narrative. Rest and locomotor-state analyses are the most direct validation modules for a Parkinson's disease sensing toolbox, and both preferentially retained beta-centered local or directional features [1-4]. By placing those panels first, the manuscript establishes that the toolbox reproduces expected physiology before asking it to solve harder event-driven problems.

The turn and cycle modules then justify the broader feature family support built into the software. In the turn module, beta-high directionality and complementary local or synchrony changes illustrate the value of discrete event alignment. In the cycle module, the retained intervals land mainly in secondary bands, but that result is still important because it demonstrates why continuous circular alignment is worth implementing in the first place. A toolbox that only recovers expected beta signatures would be useful but limited; a toolbox that also reveals sustained, phase-locked, non-beta intervals in the same anatomy-aware framework is substantially more flexible.

The study also highlights a practical principle for translational sensing workflows: support breadth should not be confused with main-text retention. The software supports periodic, aperiodic, burst, coherence, wPLI, ciPLV, TRGC, and PSI features, but the manuscript retains only the outputs that satisfy the dual-model criterion and best answer each module's question. This makes the paper more conservative while still allowing the Supplementary figures and tables to document the toolbox's full metric support.

Several limitations should be stated explicitly. The manuscript is derived from one cohort and its current summary outputs, so final patient descriptors, regulatory metadata, and funding statements still need to be reconciled against the locked clinical source tables. The cycle module currently demonstrates interval discovery most strongly in secondary bands, which means its role in the paper is extension rather than canonical validation. Finally, the copied PDFs in the submission figure folders are source panels rather than final composed figures; journal-ready assembly will still require manual layout, labeling, and typography refinement. [[COMMENT: Add the final cohort size, ethics approval, and funding statements after reconciling the clinical source workbook and the submission tracker.]]

## Methods

### Study design and cohort framing

This Brain Stimulation submission was designed as a toolbox-focused analysis of chronic STN+SNr sensing data in Parkinson's disease. Participants fulfilled clinically established Parkinson's disease criteria and contributed recordings across one or more of the supported paradigms: medication contrast, standing-to-walking contrast, turn-aligned analysis, and gait-cycle analysis [12-14]. [[COMMENT: Insert the final ethics approval identifier, consent statement wording, and the exact analyzable cohort counts for each module.]]

### Anatomy-guided localization and pair definition

Electrode localization and contact-to-region assignment followed an anatomy-guided workflow designed to preserve comparability across modules. Postoperative localization was aligned to atlas space with Lead-DBS and manual refinement as needed, and adjacent bipolar channels were assigned to STN, SNr, or intermediate locations using the paper workspace localization tables [5,6,15,16]. These region assignments were then used to define STN-local pairs, SNr-local pairs, undirected SNr-STN synchrony edges, and ordered SNr->STN directional edges.

### Toolbox preprocessing and feature extraction

The toolbox organizes preprocessing as a reproducible chain from import through export. Input recordings are transformed into bipolar adjacent-contact pairs, filtered, quality-controlled, and aligned to the annotation timeline before feature extraction. Feature families include periodic and aperiodic spectral summaries, burst rate/duration/occupation measures, undirected connectivity (coherence, wPLI, ciPLV, and related supplementary families), and directional connectivity (TRGC and PSI) [7-11,17,18]. The software is written in Python and the current paper workflow is implemented in the `paper/pd` workspace of the LFP-TensorPipe repository. [[COMMENT: Confirm the final chronic sensing hardware, recorder, sampling rate, and synchronization wording for this PD cohort.]]

### Paradigm-specific alignment modes

The same software supports four complementary alignment modes. The medication module uses condition contrast between Off and On recordings. The state module aligns blocks or annotations such as Stand and Walk. The turn module aligns windows around discrete phases including Pre, Onset, Offset, and Post. The cycle module aligns continuous gait-cycle traces after shifting ipsilateral activity into the contralateral frame, enabling a unified circular phase analysis. This shared alignment logic is one of the reasons the toolbox can be used across paradigm types without changing downstream export structure.

### Group-level scalar modeling for rest, motor, and turn analyses

Scalar analyses for medication, state, and turn panels follow the current `paper/pd` R workflow. Single-channel families such as periodic, aperiodic, raw power, and burst are modeled using region-aware linear mixed-effects structures, whereas connectivity families omit region interaction terms because the exported table already encodes a single edge class per metric. In the current codebase, each selected scalar workbook is fit with both `lmer` and `rlmer`, and main-text panels are retained only when both model classes support the effect. Estimated marginal means and Tukey-adjusted pairwise contrasts are exported next to the source workbooks for figure assembly. The figure display itself prioritizes the `lmer_model` PDF, while robust support determines main-text retention.

### Continuous cycle interval modeling

Cycle analysis follows the interval-specific workflow currently implemented in the PD paper workspace. Shifted cycle traces are first exported from the normalized summary tables, restricted to the contra-aligned frame, circularly interpolated to fill missing points, smoothed using circular LOWESS, and downsampled with a 5-point block mean to a 100-bin cycle axis. Pointwise mixed-effects models are then fit separately for each metric-band-region combination using `lmer` as the primary model and `rlmer` as a robustness layer. For local metrics, the pointwise model retains region as a fixed factor; for connectivity metrics, the pointwise model tests the exported edge directly. Contiguous runs of at least five significant bins are merged into retained cycle intervals, with wrap-around continuity enforced across the 100-to-0 cycle boundary. These interval outputs provide the basis for the cycle panels and the full interval catalogue.

### Figure assembly and plotting grammar

The submission build does not compose the final multi-panel figures directly. Instead, it copies the selected source PDFs into figure-specific folders, writes a panel manifest as YAML, and creates layout specifications for the remaining software, workflow, or export panels. For scalar boxplots, the plotting grammar is fixed across the manuscript: Tukey whiskers, black error bars for model 95% confidence intervals, black solid connected points for model means, gray dashed connected points for sample means, and significance stars at `<0.05`, `<0.01`, and `<0.001`.

### Data and code availability

Code supporting the paper workflow is available in the LFP-TensorPipe repository. Submission-ready figure manifests, table workbooks, and manuscript drafts are generated from the repository-local builder introduced for this Brain Stimulation submission. [[COMMENT: Insert the final data-availability statement, public DOI if applicable, and the corresponding access conditions for protected data.]]

## Acknowledgements

[[COMMENT: Add acknowledgements, contributors outside authorship, and technical support statements.]]

## Funding

[[COMMENT: Add all grant numbers and funding bodies in the final Brain Stimulation format.]]

## Declaration of competing interest

[[COMMENT: Confirm conflict-of-interest statements for every author.]]

## Figure Legends

### Figure 1. Cohort and toolbox overview

Overview of the cohort and software architecture. The figure introduces the shared STN+SNr cohort, anatomy-guided pair definitions, the master workflow from import to export, and the four supported alignment modes used throughout the manuscript.

### Figure 2. Rest module validation and resting network extension

The rest module reproduces canonical STN beta-low physiology while extending the same workflow to SNr and SNr-STN features. Main-text panels prioritize beta-low periodic and burst summaries, with secondary-band synchrony and directional panels retained when both standard and robust mixed-effects models support the result.

### Figure 3. State module validation and directional extension

The state module captures standing-to-walking effects using beta-centered dynamic local features and nigro-subthalamic directional coupling. Beta-low burst summaries provide the validation anchors and SNr->STN PSI serves as the retained directional extension.

### Figure 4. Event-driven modules: turn-aligned and cycle-resolved discovery

The event-driven modules demonstrate both discrete and continuous alignment logic. Turn panels emphasize beta-high directional and complementary local or synchrony features, while cycle panels highlight retained local and network intervals after continuous phase alignment.

## Supplementary Figure Summaries

### Supplementary Figure S1

Full localization and contact coverage atlas.

### Supplementary Figure S2

Rest connectivity atlas across supported synchrony and directional families.

### Supplementary Figure S3

Motor connectivity atlas across supported synchrony and directional families.

### Supplementary Figure S4

Turn connectivity atlas across supported synchrony and directional families.

### Supplementary Figure S5

Cycle local atlas for periodic, aperiodic, and optional raw-power traces.

### Supplementary Figure S6

Cycle connectivity atlas across coherence, wPLI, ciPLV, PLI, PLV, PSI, and TRGC families.

### Supplementary Figure S7

App configuration and export atlas.

## References

{references_block}
"""


def _manuscript_files() -> dict[str, str]:
    return {
        "title_page.md": _title_page_text(),
        "highlights.md": _highlights_text(),
        "main_manuscript.md": _main_manuscript_text(),
    }


def _set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def _add_text_with_comments(document: Document, paragraph: Any, text: str) -> None:
    for chunk, comment in extract_comment_segments(text):
        if not chunk and comment is None:
            continue
        run = paragraph.add_run(chunk)
        if comment is not None:
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            comment_obj = document.comments.add_comment(
                comment,
                author=COMMENT_AUTHOR,
                initials=COMMENT_INITIALS,
            )
            run.mark_comment_range(run, comment_obj.comment_id)


def markdown_to_docx(markdown_text: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_doc_defaults(doc)
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            _add_text_with_comments(doc, paragraph, line[2:].strip())
            index += 1
            continue
        if line.startswith("## "):
            paragraph = doc.add_heading(level=1)
            _add_text_with_comments(doc, paragraph, line[3:].strip())
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_heading(level=2)
            _add_text_with_comments(doc, paragraph, line[4:].strip())
            index += 1
            continue
        if line.startswith("- "):
            while index < len(lines) and lines[index].startswith("- "):
                paragraph = doc.add_paragraph(style="List Bullet")
                _add_text_with_comments(doc, paragraph, lines[index][2:].strip())
                index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index]):
                paragraph = doc.add_paragraph(style="List Number")
                text = re.sub(r"^\d+\.\s+", "", lines[index]).strip()
                _add_text_with_comments(doc, paragraph, text)
                index += 1
            continue
        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines) and lines[index].strip():
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = doc.add_paragraph()
        _add_text_with_comments(doc, paragraph, " ".join(paragraph_lines))
    doc.save(out_path)


def write_manuscript_files(project_root: Path, figure_manifest: dict[str, Any]) -> None:
    root = _manuscript_root(project_root)
    root.mkdir(parents=True, exist_ok=True)
    markdown_files = _manuscript_files()
    for filename, content in markdown_files.items():
        path = root / filename
        path.write_text(content, encoding="utf-8")
        markdown_to_docx(content, path.with_suffix(".docx"))
    manifest = {
        "reference_methods_doc": DEFAULT_REFERENCE_METHODS_DOC.as_posix(),
        "summary_table_root": summary_table_root(project_root).as_posix(),
        "stats_scripts": [
            "/Users/mojackhu/Github/LFP-TensorPipe/paper/pd/stats/run_scalar.R",
            "/Users/mojackhu/Github/LFP-TensorPipe/paper/pd/stats/run_cycle_interval_fit.R",
            "/Users/mojackhu/Github/LFP-TensorPipe/paper/pd/stats/run_cycle_interval_postprocess.R",
        ],
        "figures": figure_manifest,
        "references_count": len(_build_reference_list()),
    }
    (root / "manuscript_sources.yml").write_text(
        yaml.safe_dump(manifest, sort_keys=False, allow_unicode=False),
        encoding="utf-8",
    )


def build_tables(project_root: Path, figure_manifest: dict[str, Any]) -> None:
    root = _figure_table_root(project_root)
    root.mkdir(parents=True, exist_ok=True)
    chosen_pdfs = {
        panel["source_path"]
        for figure in figure_manifest.values()
        for panel in figure
        if panel["status"] == "copied" and str(panel["source_path"]).endswith(".pdf")
    }
    tables = {
        "Table1": build_table1(project_root),
        "Table2": build_table2(project_root),
        "Table3": build_table3(project_root),
        "TableS1": build_table_s1(project_root),
        "TableS2": build_table_s2(project_root),
        "TableS3": build_table_s3(project_root),
        "TableS4": build_table_s4(project_root, chosen_pdfs),
        "TableS5": build_table_s5(project_root),
        "TableS6": build_table_s6(project_root),
    }
    header_notes = table_header_comments()
    for name, frame in tables.items():
        _write_excel_table(
            frame,
            root / f"{name}.xlsx",
            sheet_name=name,
            header_comments=header_notes.get(name),
        )


def run(project_root: str | Path | None = None) -> None:
    project = resolve_project_root(project_root)
    figure_manifest = build_figures(project)
    build_tables(project, figure_manifest)
    write_manuscript_files(project, figure_manifest)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--project-root",
        default=None,
        help="PD project root. Defaults to paper.pd.specs.DEFAULT_PROJECT_ROOT.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    run(args.project_root)


if __name__ == "__main__":
    main()
